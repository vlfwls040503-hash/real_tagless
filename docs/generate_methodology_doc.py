"""
방법론 세팅 / 필요 데이터 / 수집 가능성 / 일정 계획 Word 문서 생성
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pathlib

OUTPUT_DIR = pathlib.Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# 제목
title = doc.add_heading('태그리스 게이트 분리 배치 최적화', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('방법론 세팅 및 연구 계획서')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2026.03.23 | 교통공학과 졸업작품')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 방법론 세팅
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. 방법론 세팅', level=1)

doc.add_heading('1.1 연구 목적', level=2)
doc.add_paragraph(
    '도시철도 개찰구에서 태그(교통카드 태핑) 이용자와 태그리스(QR/생체인식 등 무접촉) 이용자가 '
    '혼재할 때, 태그리스 전용 게이트를 분리 배치함으로써 전체 보행자 통행비용을 절감할 수 있는지 '
    '시뮬레이션 기반으로 분석한다.'
)
doc.add_paragraph(
    '분석 대상은 서울 지하철 성수역 2F 대합실(서쪽 50m × 25m, 게이트 7개)이며, '
    '태그리스 비율 변화(0~50%)와 전용 게이트 배치 위치에 따른 시나리오를 비교한다.'
)

doc.add_heading('1.2 시뮬레이션 프레임워크', level=2)

# 프레임워크 표
table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['구분', '내용']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows_data = [
    ('시뮬레이션 도구', 'JuPedSim (Python, 오픈소스 보행자 미시 시뮬레이터)'),
    ('보행 모델', 'Collision-Free Speed Model V2 (CFSM V2)\n'
     '- Tordeux et al. (2015), 속도 기반 모델\n'
     '- 기하구조 하드 바운더리 처리 (벽 관통 방지)\n'
     '- 좁은 게이트 통로(0.55m) 안정적 통과'),
    ('게이트 선택 모델', 'Gao et al. (2019) LRP 모델\n'
     '- 효용함수: V = ω_wait·대기시간 + ω_walk·보행시간\n'
     '- 3단계 재선택 (3.0m → 1.7m → 1.0m)\n'
     '- 보행자 성격 3유형 (adventurous/conserved/mild)'),
    ('서비스 시간 모델', '태그: Lognormal (평균 2.0s, σ=0.5, 범위 0.8~3.7s)\n'
     '태그리스: 0s (무정지 통과, 우이신설선 실측으로 보정 예정)'),
]
for i, (col1, col2) in enumerate(rows_data):
    table.rows[i+1].cells[0].text = col1
    table.rows[i+1].cells[1].text = col2

doc.add_paragraph()

doc.add_heading('1.3 시뮬레이션 시나리오 설계', level=2)

doc.add_paragraph('실험은 두 축으로 구성된다:')

p = doc.add_paragraph()
run = p.add_run('(1) 태그리스 비율 변화: ')
run.bold = True
p.add_run('0%, 10%, 20%, 30%, 40%, 50%')

p = doc.add_paragraph()
run = p.add_run('(2) 게이트 배치 전략: ')
run.bold = True

doc.add_paragraph('A. 기본(혼용): 모든 게이트에서 태그+태그리스 혼용', style='List Bullet')
doc.add_paragraph('B. 중앙 분리: 중앙 1~2개 게이트를 태그리스 전용으로 지정', style='List Bullet')
doc.add_paragraph('C. 양단 분리: 양쪽 끝 게이트를 태그리스 전용으로 지정', style='List Bullet')
doc.add_paragraph('D. 편측 분리: 한쪽 끝에 집중 배치', style='List Bullet')
doc.add_paragraph('E. 비율 적응형: 태그리스 비율에 비례하여 전용 게이트 수 조정', style='List Bullet')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('총 시나리오 수: ')
run.bold = True
p.add_run('6(비율) × 5(배치) = 30개, 각 시나리오 10회 반복(시드 변경) → 총 300회 시뮬레이션')

doc.add_heading('1.4 평가 지표', level=2)

table = doc.add_table(rows=6, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['지표', '정의', '출처']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

metrics = [
    ('평균 통행시간', '스폰~출구 도달 시간의 평균', '본 연구'),
    ('게이트 균형도 (MD)', 'Σ|p_j - 1/N| / N × 100%', 'Gao et al. (2019)'),
    ('평균 서비스시간', '게이트 진입~통과 완료 시간', 'Gao et al. (2019)'),
    ('최대 대기열 길이', '게이트당 피크 대기 인원', '본 연구'),
    ('통과율', '생성 대비 시뮬레이션 내 통과 비율', '본 연구'),
]
for i, (m, d, s) in enumerate(metrics):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = d
    table.rows[i+1].cells[2].text = s

doc.add_heading('1.5 보행 모델 선정 근거', level=2)

doc.add_paragraph(
    'JuPedSim에서 제공하는 보행 모델 중 CFSM V2를 선정한 근거는 다음과 같다:'
)
doc.add_paragraph(
    '(1) 기하구조 하드 바운더리: 개찰구 배리어 벽 두께가 0.30m로 매우 얇아, '
    '힘 기반 모델(SFM, GCFM)에서는 군집 압력으로 에이전트가 벽을 관통하는 문제가 발생. '
    'CFSM V2는 속도 기반으로 기하구조를 하드 바운더리로 처리하여 이를 방지.',
    style='List Bullet'
)
doc.add_paragraph(
    '(2) 좁은 통로 안정성: 게이트 통로 폭 0.55m에서 보행자 반경 0.225m 기준, '
    '클리어런스 0.05m/측. CFSM V2의 속도 조절 메커니즘이 안정적 통과를 보장.',
    style='List Bullet'
)
doc.add_paragraph(
    '(3) 대기열 행태: CFSM V2 자체는 큐 형성이 약하지만, '
    'Leader-Follower 모델(앞사람 추종 + 속도 제어)과 게이트 점유 판정을 결합하여 '
    '현실적 대기열 행태를 구현.',
    style='List Bullet'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. 방법론에 필요한 데이터
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. 방법론에 필요한 데이터', level=1)

doc.add_heading('2.1 시뮬레이션 입력 파라미터', level=2)

table = doc.add_table(rows=12, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['파라미터', '현재 값', '출처', '실측 필요 여부']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

params = [
    ('보행 희망속도 (평균)', '1.34 m/s', 'Weidmann (1993)', '△ 한국 검증'),
    ('보행 희망속도 (표준편차)', '0.26 m/s', 'Weidmann (1993)', '△ 한국 검증'),
    ('태그 서비스시간 (평균)', '2.0 s', 'Gao et al. (2019)', '○ 한국 실측 필요'),
    ('태그 서비스시간 (분포)', 'Lognormal (σ=0.5)', 'Gao et al. (2019)', '○ 한국 실측 필요'),
    ('태그리스 서비스시간', '0 s (무정지)', '본 연구 가정', '◎ 반드시 실측'),
    ('게이트 통과속도', '0.65 m/s', 'Gao et al. (2019)', '○ 한국 실측 필요'),
    ('카드 태핑시간', '1.1 s', 'Gao et al. (2019)', '△ 한국 검증'),
    ('열차 도착 간격', '180 s', '2호선 시간표', '× 공개 데이터'),
    ('1회 하차 인원', '40명 (서쪽)', '본 연구 가정', '○ 현장 관측 필요'),
    ('태그리스 이용 비율', '0~50% (시나리오)', '-', '△ 우이신설선 관측'),
    ('보행자 성격 비율', '1:1:1', 'Gao et al. (2019) 가정', '△ 검증 어려움'),
]
for i, (p1, p2, p3, p4) in enumerate(params):
    table.rows[i+1].cells[0].text = p1
    table.rows[i+1].cells[1].text = p2
    table.rows[i+1].cells[2].text = p3
    table.rows[i+1].cells[3].text = p4

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('범례: ')
run.bold = True
p.add_run('◎ 핵심(반드시 실측)  ○ 중요(실측 권장)  △ 보조(검증 수준)  × 불필요')

doc.add_heading('2.2 기하구조 데이터', level=2)

doc.add_paragraph('대합실 도면: 성수역 서쪽 대합실 50m × 25m', style='List Bullet')
doc.add_paragraph('게이트 배치: 7개, 통로폭 0.55m, 하우징폭 0.30m', style='List Bullet')
doc.add_paragraph('계단 위치: 상부(y=15~18m), 하부(y=8~11m)', style='List Bullet')
doc.add_paragraph('출구 위치: 상부(x=26~29m, y=24m), 하부(x=26~29m, y=3m)', style='List Bullet')
doc.add_paragraph('비통행 구조물: 우측 상/하 직사각형', style='List Bullet')

doc.add_heading('2.3 검증용 데이터', level=2)

doc.add_paragraph('Gao et al. (2019) Table 2~3: 게이트별 통과 분포, MD, 서비스시간 (베이징 지하철)', style='List Bullet')
doc.add_paragraph('우이신설선 현장 촬영: 태그리스 게이트 실제 이용 패턴', style='List Bullet')
doc.add_paragraph('성수역 현장 관측: 하차 인원, 동선, 피크시간대 확인', style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. 수집 가능한 데이터 vs 어려운 데이터
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. 데이터 수집 가능성 평가', level=1)

doc.add_heading('3.1 수집 가능한 데이터 (난이도 ★☆☆)', level=2)

table = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['데이터', '수집 방법', '비고']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

easy_data = [
    ('성수역 대합실 도면', 'AnyLogic 도면 + 현장 실측', '이미 확보 완료'),
    ('게이트 치수', '현장 줄자 측정', '이미 확보 완료'),
    ('열차 시간표/배차', '서울교통공사 공개 데이터', '공개 데이터'),
    ('태그 서비스시간\n(한국 실측)', '우이신설선 영상 촬영 → Kinovea', '촬영 1~2일'),
    ('태그리스 서비스시간', '우이신설선 태그리스 게이트 촬영', '촬영 1~2일'),
    ('게이트 전방 보행속도', '영상 + Kinovea/트래킹', '동시 수집 가능'),
]
for i, (d, m, n) in enumerate(easy_data):
    table.rows[i+1].cells[0].text = d
    table.rows[i+1].cells[1].text = m
    table.rows[i+1].cells[2].text = n

doc.add_heading('3.2 수집 가능하지만 노력 필요 (난이도 ★★☆)', level=2)

table = doc.add_table(rows=6, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['데이터', '수집 방법', '난이도 요인']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

mid_data = [
    ('성수역 하차 인원', '피크시간 현장 카운팅 (30분 × 2회)', '현장 인력 필요'),
    ('게이트별 대기열 길이', '영상 촬영 + 수동 카운팅', '시야각 확보 어려움'),
    ('태그리스 이용 비율', '우이신설선 게이트별 통과 유형 관측', '표본 수 확보'),
    ('보행자 경로 선택 패턴', '영상 분석 (수동 또는 반자동)', '분석 시간 소요'),
    ('피크/비피크 도착률 분포', '시간대별 카운팅 (1시간 이상)', '장시간 촬영 필요'),
]
for i, (d, m, n) in enumerate(mid_data):
    table.rows[i+1].cells[0].text = d
    table.rows[i+1].cells[1].text = m
    table.rows[i+1].cells[2].text = n

doc.add_heading('3.3 수집 어려운 데이터 (난이도 ★★★)', level=2)

table = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['데이터', '어려운 이유', '대안']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

hard_data = [
    ('보행자 성격 유형 비율\n(adventurous/conserved/mild)',
     '설문 없이 관측 불가.\n행동 기반 추정도 표본 부족.',
     'Gao (2019) 1:1:1 가정 유지\n+ 민감도 분석'),
    ('VOT(시간가치) 가중치',
     '개인별 시간가치는 설문+SP 필요.\n졸업작품 스케일에서 비현실적.',
     'Gao (2019) 가중치 사용\n+ 민감도 분석'),
    ('게이트 내부 보행 궤적\n(미시적 이동경로)',
     'CCTV 접근 불가.\n스마트폰 촬영으로는 해상도 부족.',
     'Gao (2019) 매크로 지표로 검증\n(분포, MD, 서비스시간)'),
    ('미래 태그리스 보급률',
     '정책 결정 사항으로 예측 불가.',
     '0~50% 시나리오 분석으로 대체'),
]
for i, (d, r, a) in enumerate(hard_data):
    table.rows[i+1].cells[0].text = d
    table.rows[i+1].cells[1].text = r
    table.rows[i+1].cells[2].text = a

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. 일정 계획
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. 일정 계획', level=1)

doc.add_paragraph(
    '프로젝트 확정일(2026-03-16) 기준, 학기 말 발표(6월 중순)까지 약 13주.'
)

# 간트 차트 대신 표로
table = doc.add_table(rows=10, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['단계', '기간', '주요 활동', '산출물']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

schedule = [
    ('1. 현장조사 준비',
     '3/24 ~ 3/30\n(1주)',
     '- 우이신설선 촬영 역 선정\n- 촬영 장비/위치 사전답사\n- 성수역 도면 최종 검증',
     '촬영 계획서'),
    ('2. 현장조사 실시',
     '3/31 ~ 4/13\n(2주)',
     '- 우이신설선 태그리스 게이트 촬영\n  (피크 2회 + 비피크 1회)\n- 성수역 하차 인원 관측\n- Kinovea 영상 분석',
     '영상 데이터\n파라미터 실측값'),
    ('3. 시뮬레이션 보정',
     '4/14 ~ 4/27\n(2주)',
     '- 실측 파라미터로 시뮬레이션 캘리브레이션\n- Gao (2019) 검증 지표 재확인\n- 보행 모델 최종 확정',
     '캘리브레이션 보고서\nv8 시뮬레이션'),
    ('4. 시나리오 실험',
     '4/28 ~ 5/11\n(2주)',
     '- 30개 시나리오 × 10회 반복 실행\n- 결과 데이터 정리/통계 처리\n- 최적 배치 전략 도출',
     '시나리오 결과 DB\n최적 배치 분석'),
    ('5. 민감도 분석',
     '5/12 ~ 5/18\n(1주)',
     '- 성격 비율 변화 민감도\n- VOT 가중치 변화 민감도\n- 태그리스 서비스시간 변화 민감도',
     '민감도 분석 결과'),
    ('6. 논문 작성',
     '5/19 ~ 6/1\n(2주)',
     '- 서론, 선행연구, 방법론 작성\n- 결과 분석 및 시사점\n- 그래프/표 정리',
     '논문 초고'),
    ('7. 발표 준비',
     '6/2 ~ 6/8\n(1주)',
     '- PPT 제작\n- 발표 연습\n- 논문 최종 수정',
     'PPT, 최종 논문'),
    ('8. 발표',
     '6/9 ~ 6/15',
     '- 졸업작품 발표',
     '발표 완료'),
    ('버퍼',
     '각 단계 사이\n(~2주 여유)',
     '- 예기치 못한 이슈 대응\n- 추가 현장조사 등',
     '-'),
]
for i, (s, p, a, o) in enumerate(schedule):
    table.rows[i+1].cells[0].text = s
    table.rows[i+1].cells[1].text = p
    table.rows[i+1].cells[2].text = a
    table.rows[i+1].cells[3].text = o

doc.add_paragraph()

doc.add_heading('4.1 크리티컬 패스', level=2)
doc.add_paragraph(
    '현장조사(2단계) → 캘리브레이션(3단계) → 시나리오 실험(4단계)이 순차적 종속 관계. '
    '현장조사 지연 시 전체 일정에 직접 영향.',
    style='List Bullet'
)
doc.add_paragraph(
    '논문 작성(6단계)의 서론/선행연구는 3~4단계와 병행 가능.',
    style='List Bullet'
)

doc.add_heading('4.2 리스크 요인', level=2)

table = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['리스크', '영향', '대응']):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

risks = [
    ('역사 내 촬영 제한',
     '파라미터 실측 불가',
     '서울교통공사 사전 문의\n불가 시 출입구 외부에서 촬영'),
    ('태그리스 게이트 운영 중단',
     '핵심 데이터 수집 불가',
     '다른 우이신설선 역 대체\n또는 Gao (2019) 파라미터로 보정'),
    ('시뮬레이션 모델 불안정',
     '300회 실험 일정 지연',
     '보행 모델 백업 (CFSM V2 확정)\n코드 안정화 선 완료'),
    ('표본 수 부족',
     '통계적 유의성 확보 실패',
     '최소 300명 관측 목표\n피크시간 집중 촬영'),
]
for i, (r, e, m) in enumerate(risks):
    table.rows[i+1].cells[0].text = r
    table.rows[i+1].cells[1].text = e
    table.rows[i+1].cells[2].text = m

# ── 저장 ──
output_path = OUTPUT_DIR / "방법론_연구계획서.docx"
doc.save(str(output_path))
print(f"저장 완료: {output_path}")
