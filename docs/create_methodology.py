"""졸업설계 방법론 보고서 - Word 문서 생성"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

title = doc.add_heading('태그리스 게이트 과도기 운영 전략 수립을 위한\n보행자 미시 시뮬레이션 방법론', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 1. 연구 배경 및 목적
# ============================================================
doc.add_heading('1. 연구 배경 및 목적', level=1)

doc.add_heading('1.1 연구 배경', level=2)
doc.add_paragraph(
    '서울교통공사를 비롯한 국내 도시철도 운영기관은 태그리스(tagless) 요금 징수 시스템의 '
    '도입을 추진하고 있다. 태그리스 시스템은 승객이 교통카드를 물리적으로 태깅하지 않고 '
    '자동으로 요금이 징수되는 방식으로, 게이트 통과 시 정지 또는 감속 없이 연속적인 '
    '보행이 가능하다.'
)
doc.add_paragraph(
    '그러나 태그리스 시스템의 전면 도입까지는 과도기가 불가피하다. 이 기간 동안 태그리스 '
    '이용자와 기존 교통카드(태그) 이용자가 혼재하며, 두 유형의 승객이 동일한 게이트 시설을 '
    '공유해야 한다. 이때 게이트를 어떻게 운영하느냐에 따라 — 태그리스 전용 게이트를 분리 '
    '배치할 것인지, 겸용 게이트로 운영할 것인지, 그 비율은 어떻게 할 것인지 — 역사 내 '
    '보행 서비스수준이 크게 달라질 수 있다.'
)
doc.add_paragraph(
    '현재까지 태그리스 과도기 게이트 운영 전략에 대한 체계적 연구는 부재한 상황이다. '
    '선행연구에서는 게이트 방향 제한(Peng et al., 2023), 게이트 배치 최적화(Lin et al., 2023), '
    '게이트 선택 행태 모델링(Haghani & Sarvi, 2017) 등이 수행되었으나, 태그리스/태그 혼합 '
    '운영 시나리오를 다룬 연구는 확인되지 않는다.'
)

doc.add_heading('1.2 연구 목적', level=2)
doc.add_paragraph(
    '본 연구의 목적은 보행자 미시 시뮬레이션을 통해 태그리스 과도기 게이트 운영 전략이 '
    '역사 내 보행 서비스수준에 미치는 영향을 정량적으로 평가하고, 태그리스 도입률에 따른 '
    '최적 운영 방안을 제시하는 것이다.'
)
doc.add_paragraph('구체적인 연구 질문은 다음과 같다.')
questions = [
    '태그리스 도입률(전체 승객 중 태그리스 이용자 비율)에 따라 게이트 운영 방식 간 '
    '보행 서비스수준 차이가 존재하는가?',
    '전용 분리 운영(태그리스 전용 + 태그 전용)과 겸용 운영(두 유형 모두 이용 가능) 중 '
    '어떤 방식이 더 효율적인가? 그 효과는 도입률에 따라 어떻게 변화하는가?',
    '주어진 도입률에서 전용/겸용 게이트의 최적 배분 비율은 무엇인가?',
]
for q in questions:
    doc.add_paragraph(q, style='List Number')

doc.add_heading('1.3 연구 가설', level=2)
hypotheses = [
    ('H1: ', '동일한 태그리스 도입률에서, 전용 분리 운영이 전면 겸용 운영 대비 평균 '
     '통행시간을 단축시킨다. (근거: 전용 분리 시 태그리스 이용자가 태그 이용자의 서비스 '
     '시간에 의한 지체를 받지 않으므로, 동선 교차 및 대기열 혼합이 감소한다.)'),
    ('H2: ', '도입률이 낮은 단계(α < 40%)에서는 겸용 운영이, 도입률이 높은 단계'
     '(α > 60%)에서는 전용 분리 운영이 더 효율적이다. (근거: 도입률이 낮을 때 '
     '전용 게이트를 분리하면 태그리스 전용 게이트의 활용률이 낮아져 전체 용량이 '
     '감소하고, 도입률이 높을 때는 겸용 게이트에서 태그 이용자의 서비스 시간이 '
     '태그리스 이용자의 흐름을 방해한다.)'),
    ('H3: ', '각 도입률 수준에서 전용/겸용 게이트의 최적 배분 비율이 존재하며, '
     '이는 도입률이 증가함에 따라 태그리스 전용 게이트 비율이 증가하는 양상을 보인다.'),
]
for label, text in hypotheses:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)

# ============================================================
# 2. 연구 범위 및 전제 조건
# ============================================================
doc.add_heading('2. 연구 범위 및 전제 조건', level=1)

doc.add_heading('2.1 대상 역사 및 구간', level=2)
doc.add_paragraph(
    '본 연구의 대상 역사는 서울 지하철 2호선 성수역이다. 성수역은 2호선 본선과 '
    '성수지선이 분기하는 환승역으로, 출퇴근 시간대 승하차 수요가 높은 역사이다.'
)
doc.add_paragraph(
    '시뮬레이션 범위는 성수역 2층(대합실) 서쪽 게이트 구간으로 한정한다. 해당 구간에는 '
    '개찰구 7대가 배치되어 있으며, 3층 승강장에서 계단/에스컬레이터를 통해 하차한 승객이 '
    '게이트를 통과하여 출구로 향하는 동선을 대상으로 한다.'
)

doc.add_heading('2.2 전제 조건 및 가정', level=2)
assumptions = [
    '시뮬레이션은 하차 방향(유료구역 → 무료구역) 단방향 흐름만을 대상으로 한다. '
    '승차 방향 흐름과의 교차 효과는 본 연구에서 고려하지 않는다.',
    '모든 보행자는 게이트를 통과하여 출구로 향하며, 대합실 내 체류(매표소 이용, '
    '안내판 확인 등)는 고려하지 않는다.',
    '게이트의 물리적 규격(통로 폭 0.55m, 본체 폭 0.30m, 본체 길이 1.50m)은 '
    '현행 서울 도시철도 표준을 적용한다.',
    '태그리스 이용자와 태그 이용자의 구분은 오직 게이트 서비스 시간(통과 소요 시간)의 '
    '차이로만 반영하며, 보행 속도 등 기타 보행 특성은 동일하다고 가정한다.',
    '보행자의 게이트 선택은 거리와 대기열 길이에 기반한 효용 극대화 모델을 따른다.',
]
for a in assumptions:
    doc.add_paragraph(a, style='List Number')

# ============================================================
# 3. 시뮬레이션 방법론 개요
# ============================================================
doc.add_heading('3. 시뮬레이션 방법론 개요', level=1)

doc.add_heading('3.1 방법론 흐름도', level=2)
doc.add_paragraph(
    '본 연구의 시뮬레이션 방법론은 다음과 같은 단계로 구성된다.'
)

flow_steps = [
    ('1단계: 입력 데이터 구축', '대상 역사의 기하구조, 수요 데이터(시간대별 승하차 인원), '
     '게이트 서비스 시간 실측/문헌값을 수집하여 시뮬레이션 입력 데이터를 구축한다.'),
    ('2단계: 기본 모델 구성', '보행자 이동 모델 선정, 도착 모델 설계, 게이트 의사결정 '
     '모델 설계, 서비스 시간 모델 설정 등 시뮬레이션의 기본 구성 요소를 설계한다.'),
    ('3단계: 모델 검증(Validation)', '기본 모델(Base case: 현행 운영)의 시뮬레이션 결과를 '
     '실측 데이터와 비교하여 모델의 타당성을 검증한다.'),
    ('4단계: 시나리오 설계 및 실행', '태그리스 도입률과 게이트 운영 방식을 조합한 '
     '시나리오를 설계하고, 각 시나리오별 시뮬레이션을 반복 실행한다.'),
    ('5단계: 결과 분석 및 최적안 도출', '시나리오별 성능 지표(MOE)를 비교 분석하여 '
     '도입률별 최적 게이트 운영 전략을 도출한다.'),
]
for title_text, desc in flow_steps:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    p.add_run('\n' + desc)

doc.add_heading('3.2 시뮬레이션 도구 선정', level=2)
doc.add_paragraph(
    '시뮬레이션 도구로는 JuPedSim(Juelich Pedestrian Simulator; Kemloh Wagoum et al., '
    '2015)을 사용한다. JuPedSim은 독일 율리히 연구센터에서 개발된 오픈소스 미시적 '
    '보행자 동역학 시뮬레이션 프레임워크이다.'
)
doc.add_paragraph(
    'JuPedSim 선정의 핵심 근거는 Collision-Free Speed Model V2(CFSM V2; Xu et al., '
    '2019) 모델의 지원이다. 지하철 개찰구 통로 폭(0.55m)은 보행자 직경(~0.45m) 대비 '
    '여유가 0.10m에 불과한 극히 좁은 환경으로, 기존 상용 도구(AnyLogic, PTV Vissim)가 '
    '채택한 Social Force Model(SFM; Helbing & Molnar, 1995)은 이러한 환경에서 '
    '진동(oscillation) 및 에이전트 겹침(overlapping) 문제가 보고되었다 (Kretz, 2015). '
    'CFSM V2는 속도 기반 모델로서 충돌을 원천적으로 방지하며, 병목(bottleneck) '
    '시나리오에서 실험 데이터와의 정량적 비교를 통해 검증된 바 있다 '
    '(S. Zhang et al., 2019; Xu et al., 2019).'
)

# ============================================================
# 4. 시뮬레이션 구성 요소 설계
# ============================================================
doc.add_heading('4. 시뮬레이션 구성 요소 설계', level=1)

doc.add_heading('4.1 보행자 이동 모델', level=2)

doc.add_paragraph(
    '보행자의 미시적 이동을 모사하기 위해 Collision-Free Speed Model V2(CFSM V2; '
    'Xu et al., 2019)를 채택한다. CFSM V2는 Tordeux, Chraibi, & Seyfried(2016)가 '
    '제안한 충돌 없는 속도 모델의 일반화 버전으로, 속도 기반의 1차 운동 방정식으로 '
    '보행자 이동을 기술한다.'
)

p = doc.add_paragraph()
p.add_run('핵심 수식:').bold = True

doc.add_paragraph(
    '각 보행자 i의 이동은 다음 속도 방정식으로 결정된다.'
)
p = doc.add_paragraph()
run = p.add_run('ẋᵢ = Vᵢ(sᵢ) · eᵢ')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '여기서 최적 속도 함수 V(s) = min{v₀, max{0, (s − l) / T}}이며, '
    'v₀는 희망 속도, l은 에이전트 직경, T는 시간 간격(time gap), s는 전방 최근접 '
    '이웃까지의 거리이다. 전방 이웃과의 거리가 에이전트 직경 이하가 되면 속도가 '
    '자동으로 0이 되어 충돌이 원천적으로 방지된다.'
)
doc.add_paragraph(
    '이동 방향 eᵢ는 목적지 방향 e₀에 이웃 보행자 및 벽면으로부터의 반발 함수 '
    'R(s) = a·exp((l−s)/D)를 합산하여 결정된다.'
)

p = doc.add_paragraph()
p.add_run('CFSM V2 선정 근거:').bold = True

cfsm_reasons = [
    'SFM 대비 좁은 통로(0.55m)에서 진동·겹침 없이 안정적 시뮬레이션 가능 (Kretz, 2015)',
    '에이전트별 파라미터(속도, 반경, time gap) 동적 제어 가능 — 대기열 행태, '
    '서비스 시간 모델 등 본 연구 고유의 행태 모델링에 필수적',
    '병목 시나리오에서 실험 데이터로 검증됨 (S. Zhang et al., 2019; Xu et al., 2019)',
    '파라미터가 물리적 관측량(보행 속도, 반경, 시간 간격)으로 구성되어 실측 대응 용이',
]
for r in cfsm_reasons:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
p.add_run('보행자 물리 파라미터:').bold = True

phys_data = [
    ['파라미터', '기호', '설정값', '출처'],
    ['보행자 반경', 'r', '0.225 m', 'Weidmann (1993)'],
    ['희망 보행 속도', 'v₀', 'N(1.34, 0.26) m/s', 'Weidmann (1993)'],
    ['시간 간격 (일반)', 'T', '1.06 s', 'Tordeux et al. (2016)'],
    ['시간 간격 (대기열)', 'T_q', '0.50 s', '본 연구 설정'],
    ['이웃 반발 강도', 'aₙ', '8.0', 'Tordeux et al. (2016)'],
    ['이웃 반발 범위', 'Dₙ', '0.1 m', 'Tordeux et al. (2016)'],
    ['벽면 반발 강도', 'ag', '5.0', 'Tordeux et al. (2016)'],
    ['벽면 반발 범위', 'Dg', '0.02 m', 'Tordeux et al. (2016)'],
]
t = doc.add_table(rows=len(phys_data), cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(phys_data):
    for j, cell_text in enumerate(row_data):
        cell = t.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_heading('4.2 기하구조 모델', level=2)
doc.add_paragraph(
    '시뮬레이션 공간은 성수역 2층 대합실 서쪽 게이트 구간으로 한정한다. 전체 대합실 '
    '형상을 정밀하게 재현하는 대신, 보행자의 게이트 접근 및 통과에 영향을 미치는 '
    '핵심 요소만을 모델링한다.'
)

geo_elements = [
    ('게이트 배리어: ', '개찰구 7대의 통로(0.55m)와 칸막이(0.30m)로 구성된 배리어. '
     '게이트 본체 길이는 1.50m이며, 보행자는 통로를 통해서만 유료구역에서 무료구역으로 '
     '이동할 수 있다.'),
    ('계단/에스컬레이터 출구: ', '3층 승강장에서 2층 대합실로 내려오는 보행자의 '
     '출발 지점. 성수역 서쪽 구간에는 2개소의 계단/에스컬레이터가 위치한다.'),
    ('출구: ', '게이트 통과 후 보행자가 역사를 빠져나가는 최종 목적지. '
     '출구 1, 4번이 해당 구간에 위치한다.'),
]
for bold_text, normal_text in geo_elements:
    p = doc.add_paragraph()
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

doc.add_paragraph(
    '대합실 외곽 벽면, 고객안내센터 등 구조물은 보행자의 게이트 접근 동선에 직접적 '
    '영향을 미치지 않으므로 모델링에서 제외한다.'
)

doc.add_heading('4.3 도착 모델', level=2)
doc.add_paragraph(
    '보행자의 도착 패턴은 열차 군집(platoon) 도착 모델로 구성한다. 도시철도의 특성상 '
    '보행자는 균일하게 도착하지 않고, 열차가 도착할 때마다 일시에 다수가 하차하여 '
    '계단을 통해 대합실로 진입한다.'
)

arr_data = [
    ['파라미터', '설정값', '근거'],
    ['열차 도착 간격', '180초', '서울 2호선 피크시 배차 간격 (~2.5~3분)'],
    ['1회 하차 인원', 'Poisson(λ=40)', '서쪽 계단 이용분 추정\n(실측 데이터로 보정 예정)'],
    ['계단 진입 분산', 'N(7.5, 3.75)초', '하차 후 계단 도달까지\n15초에 걸쳐 분산 도착'],
]
ta = doc.add_table(rows=len(arr_data), cols=3)
ta.style = 'Table Grid'
ta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(arr_data):
    for j, cell_text in enumerate(row_data):
        cell = ta.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '각 보행자에게는 도입률(α)에 따라 태그리스 또는 태그 이용자 유형이 확률적으로 배정된다. '
    '예를 들어 α = 40%인 시나리오에서는 각 보행자가 40% 확률로 태그리스, 60% 확률로 '
    '태그 이용자로 배정된다.'
)

doc.add_heading('4.4 게이트 서비스 시간 모델', level=2)
doc.add_paragraph(
    '게이트 서비스 시간은 보행자가 게이트 구간에 진입하여 통과를 완료하기까지의 '
    '소요 시간으로, 이용자 유형(태그리스/태그)에 따라 차별적으로 설정한다.'
)

svc_data = [
    ['이용자 유형', '서비스 시간 분포', '범위', '설명'],
    ['태그 (교통카드)', 'Lognormal\n(μ=0.35, σ=0.35)', '0.8 ~ 5.0초',
     '카드 접촉 + 인식 + 플랩 개방\n대부분 1~2초, 소수 지체'],
    ['태그리스', '≈ 0초', '0 ~ 0.3초',
     '무정지 통과\n자동 인식으로 감속 불필요'],
]
ts = doc.add_table(rows=len(svc_data), cols=4)
ts.style = 'Table Grid'
ts.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(svc_data):
    for j, cell_text in enumerate(row_data):
        cell = ts.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '태그 이용자의 서비스 시간을 로그정규분포로 모델링한 근거는, 서비스 시간이 '
    '양(+)의 값만 취하며 오른쪽 꼬리가 긴 비대칭 분포를 보이기 때문이다. '
    '대부분의 이용자는 1~2초 내에 빠르게 통과하나, 카드 인식 실패, 잔액 부족, '
    '카드 꺼내기 지연 등으로 소수가 장시간 지체하는 현상을 반영한다.'
)
doc.add_paragraph(
    '서비스 시간 파라미터의 최종 설정값은 현장 실측 데이터 수집 후 보정할 예정이다.'
)

doc.add_heading('4.5 게이트 선택 의사결정 모델', level=2)
doc.add_paragraph(
    '보행자의 게이트 선택은 이산선택모형(Discrete Choice Model) 기반의 로짓 모델로 '
    '구현한다. Haghani & Sarvi(2017)의 연구에 기반하여, 보행자는 각 게이트에 대한 '
    '효용(utility)을 평가하고 확률적으로 게이트를 선택한다.'
)

p = doc.add_paragraph()
p.add_run('효용함수:').bold = True
p = doc.add_paragraph()
run = p.add_run('U(i) = β_dist × dᵢ + β_queue × qᵢ')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('dᵢ : 보행자 현재 위치에서 게이트 i까지의 유클리드 거리 (m)', style='List Bullet')
doc.add_paragraph('qᵢ : 게이트 i에 현재 배정된 대기 인원 (명)', style='List Bullet')
doc.add_paragraph('β_dist = -0.25 : 거리 민감도 (Haghani & Sarvi, 2017 범위 내)', style='List Bullet')
doc.add_paragraph('β_queue = -0.30 : 대기열 민감도 (Haghani & Sarvi, 2017 범위 내)', style='List Bullet')

doc.add_paragraph(
    '선택 확률은 다항 로짓 모델(Multinomial Logit Model)에 따라 '
    'P(i) = exp(U(i)) / Σ exp(U(j))로 산출한다.'
)

p = doc.add_paragraph()
p.add_run('이용 가능 게이트 필터링:').bold = True
doc.add_paragraph(
    '보행자의 이용자 유형(태그리스/태그)에 따라 이용 가능한 게이트 선택지가 제한된다. '
    '이는 본 연구의 핵심 메커니즘으로, 전용/겸용 운영 방식에 따른 효과 차이를 발생시키는 '
    '요인이다.'
)
filter_rules = [
    '태그리스 이용자: 태그리스 전용 게이트 + 겸용 게이트만 선택 가능',
    '태그 이용자: 태그 전용 게이트 + 겸용 게이트만 선택 가능',
    '겸용 게이트: 두 유형 모두 이용 가능 (서비스 시간은 이용자 유형에 따름)',
]
for rule in filter_rules:
    doc.add_paragraph(rule, style='List Bullet')

p = doc.add_paragraph()
p.add_run('다단계 의사결정 (핑퐁 효과 방지):').bold = True
doc.add_paragraph(
    '보행자가 게이트를 선택한 후 이동 중에 주변 상황이 변화하면 경로를 재평가할 수 있다. '
    '그러나 과도한 경로 변경은 비현실적인 핑퐁 효과를 유발하므로, 다음 세 가지 '
    '메커니즘으로 제어한다.'
)

decision_data = [
    ['메커니즘', '파라미터', '설명'],
    ['관성 (Inertia)', 'C_switch = 1.5',
     '새 게이트의 효용이 현재 게이트 효용 +\n전환 비용을 초과해야 경로 변경'],
    ['Lock-in', 'd_lock = 3.0 m',
     '게이트까지 3m 이내 접근 시\n경로 변경 비활성화'],
    ['재평가 주기', 'Δt = 3.0 s',
     '경로 재평가를 3초 간격으로 제한'],
]
td = doc.add_table(rows=len(decision_data), cols=3)
td.style = 'Table Grid'
td.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(decision_data):
    for j, cell_text in enumerate(row_data):
        cell = td.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ============================================================
# 5. 시나리오 설계
# ============================================================
doc.add_heading('5. 시나리오 설계', level=1)

doc.add_heading('5.1 독립변수', level=2)

p = doc.add_paragraph()
p.add_run('(1) 태그리스 도입률 (α):').bold = True
doc.add_paragraph(
    '전체 승객 중 태그리스 시스템을 이용하는 승객의 비율. '
    '20%, 40%, 60%, 80%의 4단계로 설정한다.'
)

p = doc.add_paragraph()
p.add_run('(2) 게이트 운영 방식:').bold = True
gate_types = [
    '전용 분리: 태그리스 전용 N대 + 태그 전용 M대 (N + M = 7)',
    '겸용 혼합: 태그리스 전용 + 태그 전용 + 겸용 게이트 조합',
    '전면 겸용: 7대 전부 겸용 (두 유형 모두 이용 가능)',
]
for gt in gate_types:
    doc.add_paragraph(gt, style='List Bullet')

doc.add_heading('5.2 시나리오 매트릭스', level=2)
doc.add_paragraph(
    '게이트 총 7대 기준, 도입률(α)과 운영 방식을 조합한 시나리오를 다음과 같이 구성한다.'
)

scenario_data = [
    ['시나리오', '도입률\n(α)', '태그리스\n전용', '태그\n전용', '겸용', '비고'],
    ['Base', '-', '0', '7', '0', '현행 운영\n(태그리스 없음)'],
    ['S1-a', '20%', '1', '6', '0', '전용 분리'],
    ['S1-b', '20%', '0', '5', '2', '겸용 혼합'],
    ['S1-c', '20%', '0', '0', '7', '전면 겸용'],
    ['S2-a', '40%', '2', '5', '0', '전용 분리'],
    ['S2-b', '40%', '1', '3', '3', '겸용 혼합'],
    ['S2-c', '40%', '0', '0', '7', '전면 겸용'],
    ['S3-a', '60%', '4', '3', '0', '전용 분리'],
    ['S3-b', '60%', '2', '2', '3', '겸용 혼합'],
    ['S3-c', '60%', '0', '0', '7', '전면 겸용'],
    ['S4-a', '80%', '5', '2', '0', '전용 분리'],
    ['S4-b', '80%', '3', '1', '3', '겸용 혼합'],
    ['S4-c', '80%', '0', '0', '7', '전면 겸용'],
]
tsc = doc.add_table(rows=len(scenario_data), cols=6)
tsc.style = 'Table Grid'
tsc.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(scenario_data):
    for j, cell_text in enumerate(row_data):
        cell = tsc.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '각 시나리오는 동일한 난수 시드(random seed)를 사용하되, 통계적 신뢰성 확보를 위해 '
    '시드를 변경하며 복수 회(N ≥ 10) 반복 실행한다.'
)

doc.add_heading('5.3 종속변수 (성능 지표, MOE)', level=2)

moe_data = [
    ['지표', '정의', '단위'],
    ['평균 통행시간', '계단 진입 ~ 출구 도달\n소요 시간의 평균', '초 (s)'],
    ['최대 대기열 길이', '게이트 앞 대기 인원의\n시간대별 최댓값', '명'],
    ['게이트별 처리량', '단위 시간당\n게이트 통과 인원', '명/분'],
    ['보행자 밀도 (LOS)', '게이트 앞 구간의\n보행자 밀도', 'm²/인\n(Fruin 기준)'],
    ['게이트 활용률', '각 게이트의 가동률\n(서비스 중 시간 / 전체 시간)', '%'],
]
tm = doc.add_table(rows=len(moe_data), cols=3)
tm.style = 'Table Grid'
tm.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(moe_data):
    for j, cell_text in enumerate(row_data):
        cell = tm.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ============================================================
# 6. 모델 검증 계획
# ============================================================
doc.add_heading('6. 모델 검증 계획', level=1)

doc.add_paragraph(
    '시뮬레이션 모델의 타당성을 확보하기 위해, Base case(현행 운영)의 시뮬레이션 결과를 '
    '현장 실측 데이터와 비교 검증한다.'
)

doc.add_heading('6.1 검증 데이터 수집 계획', level=2)
validation_items = [
    '게이트별 통과 인원 (시간대별): 피크 시간대 15분 단위 게이트별 통과 인원 관측',
    '게이트 서비스 시간: 카드 태핑부터 통과 완료까지 소요 시간 실측 (최소 100건)',
    '대기열 길이: 게이트 앞 대기 인원의 시간대별 관측',
    '보행자 밀도: 게이트 전방 구간의 보행자 밀도 영상 분석 (가능한 경우)',
]
for item in validation_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 검증 기준', level=2)
doc.add_paragraph(
    '시뮬레이션 결과와 실측 데이터의 적합도는 다음 지표로 평가한다.'
)
validation_criteria = [
    'GEH 통계량: 교통 시뮬레이션 분야에서 널리 사용되는 적합도 지표로, '
    'GEH < 5이면 적합한 것으로 판단한다.',
    'RMSE (Root Mean Square Error): 시뮬레이션과 실측 간 통과 인원의 '
    '평균 제곱근 오차.',
    '시각적 비교: 대기열 형성 패턴, 보행자 흐름 양상 등을 정성적으로 비교한다.',
]
for c in validation_criteria:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# 7. 연구 한계 및 향후 과제
# ============================================================
doc.add_heading('7. 연구 한계 및 향후 과제', level=1)
limitations = [
    '본 연구는 성수역 서쪽 게이트 7대 구간만을 대상으로 하며, 역사 전체 또는 '
    '타 역사로의 일반화에는 추가 연구가 필요하다.',
    '하차 방향 단방향 흐름만을 고려하였으며, 승차 방향 흐름과의 교차 효과는 '
    '향후 양방향 시뮬레이션을 통해 분석할 필요가 있다.',
    '태그리스 시스템의 인식 실패, 오탐 등 시스템 오류 상황은 고려하지 않았다.',
    'CFSM V2 모델은 병목 시나리오에서 검증되었으나, 지하철 게이트에 직접 적용한 '
    '선행 사례가 없어 모델 적합성에 대한 추가 검증이 필요하다.',
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

# ============================================================
# 참고문헌
# ============================================================
doc.add_heading('참고문헌', level=1)
refs = [
    'Haghani, M. & Sarvi, M. (2017). Stated and revealed exit choices of pedestrian crowd evacuees. Transportation Research Part B: Methodological, 106, 410-427.',
    'Helbing, D. & Molnar, P. (1995). Social Force Model for Pedestrian Dynamics. Physical Review E, 51(5), 4282-4286.',
    'Kemloh Wagoum, A.U., Chraibi, M., & Zhang, J. (2015). JuPedSim: an open framework for simulating and analyzing the dynamics of pedestrians. 3rd Conference of Transportation Research Group of India.',
    'Kretz, T. (2015). On Oscillations in the Social Force Model. Physica A, 438, 272-285.',
    'Lin, X., Cheng, L., Zhang, S., & Wang, Q. (2023). Simulating the Effects of Gate Machines on Crowd Traffic Based on the Modified Social Force Model. Mathematics, 11(3), 780.',
    'Peng, J., Wei, Z., Wang, S., & Qiu, S. (2023). Toward Dynamic Regulation of Bidirectional Automatic Fare Gates. Simulation Modelling Practice and Theory, 124.',
    'Peng, J., Wei, Z., Li, J., Guo, X., & Wang, S. (2024). Passenger Flow Bottleneck Decongestion in Subway Stations. SIMULATION (SAGE).',
    'Rzezonka, J., Chraibi, M., Seyfried, A., Hein, B., & Schadschneider, A. (2022). An attempt to distinguish physical and socio-psychological influences on pedestrian bottleneck. Royal Society Open Science, 9(6), 211822.',
    'Seyfried, A., Steffen, B., Klingsch, W., & Boltes, M. (2005). The fundamental diagram of pedestrian movement revisited. Journal of Statistical Mechanics, P10002.',
    'Tordeux, A., Chraibi, M., & Seyfried, A. (2016). Collision-Free Speed Model for Pedestrian Dynamics. In: Traffic and Granular Flow \'15, pp. 225-232. Springer.',
    'Weidmann, U. (1993). Transporttechnik der Fussgaenger. Schriftenreihe des IVT Nr. 90, ETH Zuerich.',
    'Xu, Q., Chraibi, M., Tordeux, A., & Zhang, J. (2019). Generalized collision-free velocity model for pedestrian dynamics. Physica A, 535, 122521.',
    'Zhang, S., Chraibi, M., Zhang, J., Li, H., & Song, W. (2019). Validation of Collision-Free Speed Model in Bottlenecks. 9th ICFSFPE, IEEE.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)

output_path = r'C:\Users\aaron\tagless\docs\방법론_보고서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
