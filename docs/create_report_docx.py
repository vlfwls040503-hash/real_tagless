"""마크다운 보고서를 Word 문서로 변환"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 제목
title = doc.add_heading('보행자 미시 시뮬레이션 모델 선정 및 파라미터 설정 보고서', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 1장
# ============================================================
doc.add_heading('1. 시뮬레이션 프레임워크 선정: JuPedSim', level=1)

doc.add_heading('1.1 JuPedSim 개요', level=2)
doc.add_paragraph(
    'JuPedSim(Juelich Pedestrian Simulator)은 독일 율리히 연구센터'
    '(Forschungszentrum Juelich) 시민안전연구소(IAS-7)에서 2010년부터 개발된 '
    '오픈소스 미시적(microscopic) 보행자 동역학 시뮬레이션 프레임워크이다 '
    '(Kemloh Wagoum et al., 2015). C++ 코어에 Python 바인딩으로 구성되어 있으며, '
    'GNU LGPLv3 라이선스로 배포된다.'
)

doc.add_heading('1.2 선정 근거', level=2)

doc.add_paragraph(
    '지하철 역사 보행자 시뮬레이션에 가장 널리 사용되는 도구는 AnyLogic과 PTV Vissim/Viswalk이다 '
    '(Peng et al., 2024; Nanjing Metro study, 2021). 두 도구 모두 내장 게이트 모듈, 게이트 통과 '
    '시간 설정, 검증된 보행자 이동 로직 등 완성도 높은 기능을 제공하며, 다수의 선행연구에서 그 '
    '유효성이 입증되었다. 그럼에도 불구하고 본 연구에서 JuPedSim을 선정한 근거는 다음과 같다.'
)

reasons = [
    ('첫째, 좁은 게이트 통로에 적합한 보행자 이동 모델.',
     ' 본 연구의 핵심 시뮬레이션 환경인 지하철 개찰구 통로 폭은 0.55m로, 보행자 직경(~0.45m) '
     '대비 여유가 0.10m에 불과하다. AnyLogic과 PTV Vissim은 Social Force Model(SFM; Helbing & '
     'Molnar, 1995)을 기반으로 하는데, SFM은 이러한 극히 좁은 통로에서 진동(oscillation), '
     '에이전트 겹침(overlapping), 출구 막힘 등 구조적 한계가 보고되었다 (Kretz, 2015). '
     'JuPedSim이 지원하는 Collision-Free Speed Model V2(CFSM V2; Xu et al., 2019)는 '
     '속도 기반 모델로서 충돌을 원천적으로 방지하며, 좁은 게이트 통로에서도 안정적인 '
     '시뮬레이션이 가능하다. 특히 CFSM은 병목(bottleneck) 시나리오에서 실험 데이터와의 '
     '정량적 비교를 통해 검증된 바 있으며 (S. Zhang et al., 2019; Xu et al., 2019), '
     '게이트 통로는 물리적으로 병목 문제와 동일하므로 검증된 모델의 새로운 적용(novel '
     'application)에 해당한다. 이것이 JuPedSim 선정의 가장 본질적인 이유이다 '
     '(상세 비교는 2장에서 기술한다).'),
    ('둘째, 에이전트별 파라미터 동적 제어.',
     ' CFSM V2 모델은 각 에이전트(보행자)마다 보행 속도, 반경, time gap 등의 파라미터를 '
     '개별적으로 설정하고, 시뮬레이션 도중에도 동적으로 변경할 수 있다. 이를 통해 게이트 '
     '대기열에서의 time gap 감소(바짝 붙어 서는 심리), 서비스 중 속도 감소, 태그리스/태그 '
     '사용자 유형별 차별화된 행태 등 본 연구에서 요구하는 세밀한 행태 모델링이 가능하다.'),
    ('셋째, 오픈소스 기반의 부수적 장점.',
     ' JuPedSim은 오픈소스(GNU LGPLv3)로서 시뮬레이션 코드와 파라미터를 완전히 공개할 수 '
     '있어 연구의 재현성 확보에 유리하다. 또한 Python API를 통해 태그리스/태그 혼합 운영 등 '
     '본 연구 고유의 의사결정 로직을 유연하게 구현할 수 있다. 다만 이는 AnyLogic(Java)이나 '
     'PTV Vissim(COM API)에서도 일정 수준 가능한 기능이므로, 부수적 장점으로 간주한다.'),
]
for bold_text, normal_text in reasons:
    p = doc.add_paragraph()
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    p.add_run(normal_text)

doc.add_heading('1.3 주요 상용 도구와의 비교', level=2)

# 비교표
table_data = [
    ['항목', 'JuPedSim', 'AnyLogic', 'PTV Vissim/Viswalk'],
    ['보행자 모델', 'CFSM V1/V2', 'Social Force Model', 'Social Force Model'],
    ['좁은 통로(0.55m)\n안정성', '높음\n(충돌 원천 방지)', '보통\n(진동·겹침 가능)', '보통\n(진동·겹침 가능)'],
    ['내장 게이트 모듈', '없음\n(코드 구현 필요)', '있음\n(fare gate 내장)', '있음\n(게이트 모듈)'],
    ['에이전트별\n파라미터 동적 제어', '지원\n(시뮬레이션 중 변경)', '제한적', '제한적'],
    ['선행연구 활용도', '낮음\n(게이트 시뮬 사례 없음)', '높음\n(다수 논문)', '높음\n(다수 논문)'],
    ['라이선스', '오픈소스 (LGPL)', '상용 (학생 PLE 제한)', '상용'],
]
table = doc.add_table(rows=len(table_data), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(table_data):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '위 비교에서 확인할 수 있듯이, AnyLogic과 PTV Vissim은 내장 게이트 모듈과 풍부한 '
    '선행연구 사례 등에서 우위를 가진다. 그러나 두 도구 모두 Social Force Model 기반으로서, '
    '개찰구 통로 폭(0.55m)과 같은 극히 좁은 환경에서의 시뮬레이션 안정성에 구조적 한계가 있다. '
    '본 연구에서는 이 좁은 게이트 통로에서의 보행자 흐름이 시뮬레이션의 핵심 대상이므로, '
    '모델의 물리적 안정성을 최우선 기준으로 JuPedSim(CFSM V2)을 선정하였다.'
)

# ============================================================
# 2장
# ============================================================
doc.add_heading('2. 보행자 이동 모델 선정: Collision-Free Speed Model V2', level=1)

doc.add_heading('2.1 Social Force Model의 한계', level=2)
doc.add_paragraph(
    'Social Force Model(SFM; Helbing & Molnar, 1995)은 보행자 시뮬레이션에서 가장 널리 '
    '사용되는 모델이나, 좁은 통로 환경에서 다음과 같은 구조적 한계가 보고되었다.'
)

sfm_limits = [
    ('진동 문제 (Oscillation).',
     ' Kretz(2015)는 SFM에서 보행자가 목적지 좌표 주위를 비현실적으로 진동하는 현상을 '
     '수학적으로 분석하였다. 일반적 파라미터(이완 시간 τ = 0.4s, 희망 속도 v₀ = 1.5 m/s) '
     '하에서 진폭이 1cm 미만이 되기까지 약 7회 진동(~2.1초)이 필요하며, 이는 좁은 게이트 '
     '통로에서 비현실적인 보행 궤적을 유발한다.'),
    ('에이전트 겹침 (Overlapping).',
     ' SFM은 힘(force) 기반 모델로서, 반발력의 크기가 충분하지 않거나 시간 스텝이 클 경우 '
     '보행자 간 물리적 겹침이 발생할 수 있다. 이를 방지하기 위해 반발력을 증가시키면 진동이 '
     '심화되는 이중 딜레마(dual dilemma)가 존재한다 (Kretz, 2015).'),
    ('좁은 출구에서의 막힘.',
     ' 출구 모서리에서 과도한 심리적 힘(psychological force)으로 인해 대피 불가능 현상이 '
     '발생할 수 있다. 지하철 게이트 통로 폭(0.55m)은 보행자 직경(~0.45m) 대비 여유가 '
     '0.10m에 불과하여, SFM의 이러한 한계가 극대화되는 환경이다.'),
]
for bold_text, normal_text in sfm_limits:
    p = doc.add_paragraph()
    run_bold = p.add_run(bold_text)
    run_bold.bold = True
    p.add_run(normal_text)

doc.add_heading('2.2 Collision-Free Speed Model (CFSM) 개요', level=2)
doc.add_paragraph(
    'CFSM은 Tordeux, Chraibi, & Seyfried(2016)가 제안한 속도(velocity) 기반 보행자 이동 '
    '모델로, SFM의 구조적 한계를 해결하기 위해 개발되었다. 이후 Xu, Chraibi, Tordeux, & '
    'Zhang(2019)에 의해 일반화된 V2 모델이 제안되었다.'
)

doc.add_heading('2.3 CFSM V2 수학적 공식', level=2)

p = doc.add_paragraph()
p.add_run('속도 방정식:').bold = True
doc.add_paragraph('각 보행자 i의 이동 속도는 다음과 같이 결정된다.')
p = doc.add_paragraph()
run = p.add_run('ẋᵢ = Vᵢ(sᵢ) · eᵢ')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('여기서 Vᵢ는 속도 함수(optimal velocity function), eᵢ는 이동 방향 함수이다.')

p = doc.add_paragraph()
p.add_run('최적 속도 함수 (Optimal Velocity Function):').bold = True
p = doc.add_paragraph()
run = p.add_run('V(s) = min{v₀, max{0, (s − l) / T}}')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

params_velocity = [
    'v₀ : 희망 속도 (desired speed, m/s)',
    'l : 에이전트 직경 (= 2 × radius, m)',
    'T : 시간 간격 (time gap, s) — 안전 차간거리 유지를 위한 최소 시간',
    's : 전방 최근접 이웃까지의 거리 (m)',
]
for param in params_velocity:
    doc.add_paragraph(param, style='List Bullet')

doc.add_paragraph(
    '이 함수의 핵심적 특성은, 전방 이웃과의 거리 s가 에이전트 직경 l 이하가 되면 속도가 '
    '자동으로 0이 되어 충돌이 원천적으로 방지된다는 점이다.'
)

p = doc.add_paragraph()
p.add_run('방향 함수 (Direction Function):').bold = True
p = doc.add_paragraph()
run = p.add_run('eᵢ = normalize(e₀ + Σⱼ Rₙ(sᵢⱼ)·nᵢⱼ + Rg(sᵢw)·nᵢw)')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

params_direction = [
    'e₀ : 목적지 방향 (desired direction)',
    'Rₙ : 이웃 보행자 반발 함수',
    'Rg : 벽/기하구조 반발 함수',
    'nᵢⱼ, nᵢw : 이웃/벽으로부터의 단위 법선 벡터',
]
for param in params_direction:
    doc.add_paragraph(param, style='List Bullet')

p = doc.add_paragraph()
p.add_run('반발 함수 (Repulsion Function):').bold = True
p = doc.add_paragraph()
run = p.add_run('R(s) = a · exp((l − s) / D)')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('a : 반발 강도 (strength)', style='List Bullet')
doc.add_paragraph('D : 반발 범위 (range) — 반발력이 유효한 거리 척도', style='List Bullet')

doc.add_heading('2.4 SFM 대비 CFSM의 구조적 차이', level=2)

compare_data = [
    ['특성', 'SFM (Helbing, 1995)', 'CFSM V2 (Xu et al., 2019)'],
    ['기본 접근', '뉴턴 역학 (F = ma)', '1차 속도 방정식'],
    ['상호작용', '반발력(force)', '방향 편향(deflection)'],
    ['충돌 처리', '사후 반발 (겹침 발생 가능)', '사전 방지 (속도 → 0)'],
    ['진동', '발생 (Kretz, 2015)', '발생하지 않음'],
    ['계산 비용', '2차 미분방정식 (높음)', '1차 방정식 (낮음)'],
    ['좁은 통로', '불안정', '안정적'],
    ['파라미터 해석', '사회적 힘 (추상적)', '물리적 관측량 (직관적)'],
]
table2 = doc.add_table(rows=len(compare_data), cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(compare_data):
    for j, cell_text in enumerate(row_data):
        cell = table2.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_heading('2.5 V1 대비 V2의 개선점', level=2)
doc.add_paragraph(
    'CFSM V2 (Xu et al., 2019)는 V1 대비 다음과 같은 개선이 이루어졌다.'
)

v2_improvements = [
    ('벽면 반발력 통합: ',
     'V1에서는 이웃 보행자 반발만 고려하였으나, V2에서는 벽/기하구조의 반발력을 모델에 '
     '통합하여 벽면 근처에서의 보행 행태를 현실적으로 반영한다.'),
    ('에이전트별 개별 파라미터: ',
     'V1은 전역(global) 파라미터만 지원하였으나, V2는 각 에이전트마다 보행 속도, 반경, '
     'time gap, 반발 강도/범위를 개별 설정하고 시뮬레이션 도중 동적으로 변경할 수 있다.'),
    ('방향 함수 개선: ',
     '보행자 방향 전환 시 더 부드러운 변화를 제공하여 비현실적 후진을 감소시켰다.'),
]
for i, (bold_text, normal_text) in enumerate(v2_improvements, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

doc.add_paragraph(
    '본 연구에서는 게이트 대기열에서의 time gap 동적 변경(대기 시 바짝 붙어 서는 심리 반영), '
    '서비스 중 속도 감소 등 에이전트별 파라미터 동적 제어가 필수적이므로 V2 모델을 채택하였다.'
)

doc.add_heading('2.6 본 연구에서의 CFSM V2 선정 근거 요약', level=2)
summary_points = [
    '게이트 통로 폭(0.55m)이 보행자 직경(~0.45m) 대비 극히 좁아 SFM의 진동·겹침 문제가 '
    '심각하게 발생하는 환경이다.',
    'CFSM V2는 속도 기반 모델로 충돌을 원천 방지하므로, 좁은 게이트 통로에서도 안정적인 '
    '시뮬레이션이 가능하다.',
    '에이전트별 파라미터 동적 제어를 통해 대기열 행태(time gap 감소), 서비스 시간 모델'
    '(속도 감소), 태그리스/태그 사용자 유형별 차별화된 행태를 구현할 수 있다.',
    '파라미터가 물리적으로 관측 가능한 양(보행 속도, 보행자 반경, 시간 간격)으로 구성되어, '
    '선행연구 실측 데이터와의 대응이 용이하다.',
]
for i, point in enumerate(summary_points, 1):
    doc.add_paragraph(f'{i}. {point}')

doc.add_heading('2.7 CFSM의 병목 시나리오 검증 사례', level=2)

doc.add_paragraph(
    '본 연구에서 CFSM V2를 지하철 게이트 시뮬레이션에 적용하는 것은, 기존에 CFSM으로 '
    '게이트를 직접 시뮬레이션한 선행연구가 없다는 점에서 새로운 시도이다. 그러나 게이트 '
    '통로(폭 0.55m, 길이 1.5m)는 물리적으로 병목(bottleneck) 문제와 본질적으로 동일하며, '
    'CFSM은 병목 시나리오에서 실험 데이터를 통해 충분히 검증되었다. 주요 검증 사례는 '
    '다음과 같다.'
)

validation_data = [
    ['논문', '검증 시나리오', '핵심 결과'],
    ['S. Zhang et al.\n(2019, IEEE)', '병목 직접 검증',
     '밀도 < 3.33 m⁻²에서\nflow-density 관계 재현'],
    ['Xu et al.\n(2019, Physica A)', '복도 + 병목',
     '동적 타원형 도입,\n기본 다이어그램 검증'],
    ['S. Zhang et al.\n(2021, CNSNS)', '병목, 궤적 분포',
     '보행 선호도 반영으로\n궤적 현실성 향상'],
    ['Rzezonka et al.\n(2022, Royal Society)', '병목 앞 복도 폭 변화',
     '물리적 원인으로\n밀도 증가 현상 설명'],
    ['Seyfried et al.\n(2005; 2010)', '병목 실험 데이터',
     'CFSM 파라미터 보정에\n사용된 기초 실험 데이터'],
]
tv = doc.add_table(rows=len(validation_data), cols=3)
tv.style = 'Table Grid'
tv.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(validation_data):
    for j, cell_text in enumerate(row_data):
        cell = tv.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '특히 S. Zhang et al.(2019)은 CFSM을 실험적 병목 데이터와 정량적으로 비교하여, '
    '밀도 3.33 m⁻² 미만의 조건에서 흐름(flow)과 밀도(density) 간의 관계를 재현할 수 '
    '있음을 확인하였다. 지하철 개찰구의 일반적 운영 조건은 이 밀도 범위에 해당하므로, '
    'CFSM의 적용 타당성이 확보된다.'
)
doc.add_paragraph(
    '또한 Xu et al.(2019)은 CFSM V2(일반화된 충돌 회피 속도 모델)를 복도 및 병목 '
    '시나리오에서 검증하여, 좁은 공간에서의 기본 다이어그램(fundamental diagram)과 '
    '군중 분포가 실험 데이터와 부합함을 확인하였다.'
)
doc.add_paragraph(
    '따라서 본 연구에서의 CFSM V2 적용은 검증되지 않은 모델의 사용이 아니라, '
    '병목 시나리오에서 검증된 모델을 물리적으로 동일한 게이트 통과 문제에 적용하는 '
    '것으로서, 방법론적 타당성을 가진다.'
)

# ============================================================
# 3장
# ============================================================
doc.add_heading('3. 핵심 파라미터 설정', level=1)

doc.add_heading('3.1 보행자 물리 파라미터', level=2)

phys_data = [
    ['파라미터', '기호', '설정값', '출처'],
    ['보행자 반경', 'r', '0.225 m', 'Weidmann (1993)'],
    ['희망 보행 속도', 'v₀', 'N(1.34, 0.26) m/s', 'Weidmann (1993)'],
    ['속도 범위', '-', '0.5 ~ 2.0 m/s (clip)', '-'],
]
t = doc.add_table(rows=len(phys_data), cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(phys_data):
    for j, cell_text in enumerate(row_data):
        cell = t.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'Weidmann(1993)은 ETH Zurich에서 발표된 보행자 교통공학 분야의 기초 문헌으로, '
    '비혼잡 상태의 자유보행속도(free-flow speed) 분포를 정규분포 N(1.34, 0.26) m/s로 '
    '제시하였다. 보행자 시뮬레이션 분야에서 가장 널리 인용되는 보행 속도 파라미터 출처이며, '
    'JuPedSim의 공식 예제에서도 이 값을 기본으로 사용한다.'
)
doc.add_paragraph(
    '보행자 반경 0.225m(직경 0.45m)는 성인 어깨폭의 절반에 해당하며, 게이트 통로 폭(0.55m) '
    '대비 양측 0.05m의 여유를 제공한다.'
)

doc.add_heading('3.2 이동 모델 파라미터', level=2)

model_data = [
    ['파라미터', '기호', '설정값', '출처'],
    ['시간 간격 (일반)', 'T', '1.06 s', 'Tordeux et al. (2016)'],
    ['시간 간격 (대기열)', 'T_q', '0.50 s', '본 연구 설정'],
    ['이웃 반발 강도', 'aₙ', '8.0', 'Tordeux et al. (2016)'],
    ['이웃 반발 범위', 'Dₙ', '0.1 m', 'Tordeux et al. (2016)'],
    ['벽면 반발 강도', 'ag', '5.0', 'Tordeux et al. (2016)'],
    ['벽면 반발 범위', 'Dg', '0.02 m', 'Tordeux et al. (2016)'],
]
t2 = doc.add_table(rows=len(model_data), cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(model_data):
    for j, cell_text in enumerate(row_data):
        cell = t2.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '시간 간격(time gap) T = 1.06s는 Tordeux et al.(2016)이 보행자 추종(follow-the-leader) '
    '실험 데이터에서 속도-간격 간의 선형 관계를 분석하여 도출한 경험적 값이다. 이는 보행자가 '
    '전방 보행자와 안전 거리를 유지하기 위한 최소 시간 간격을 의미한다.'
)
doc.add_paragraph(
    '대기열 시간 간격 T_q = 0.50s는 게이트 앞 대기열에서 보행자가 일반 보행 시보다 전방 '
    '보행자에 바짝 붙어 서는 심리를 반영한 본 연구의 설정값이다.'
)
doc.add_paragraph(
    '이웃 반발 강도 aₙ = 8.0, 범위 Dₙ = 0.1m는 반발 함수 R(s) = 8.0 × exp((l − s) / 0.1)로, '
    '이웃 보행자가 10cm 이내로 접근할 때 반발이 급격히 증가하는 특성을 가진다. '
    'Tordeux et al.(2016) 원논문에서 실험 데이터 기반으로 캘리브레이션된 값이다.'
)
doc.add_paragraph(
    '벽면 반발 강도 ag = 5.0, 범위 Dg = 0.02m는 벽면 2cm 이내에서 강한 반발이 작용하도록 '
    '설정된 값으로, 보행자가 벽면에 밀착하지 않는 현실적 행태를 반영한다.'
)

doc.add_heading('3.3 게이트 선택 의사결정 파라미터', level=2)

gate_data = [
    ['파라미터', '기호', '설정값', '출처'],
    ['거리 민감도', 'β_dist', '-0.25', 'Haghani & Sarvi (2017) 범위 내'],
    ['대기열 민감도', 'β_queue', '-0.30', 'Haghani & Sarvi (2017) 범위 내'],
    ['시야 반경', 'R_v', '8.0 m', '본 연구 설정'],
    ['전환 비용 (관성)', 'C_switch', '1.5', '본 연구 설정'],
    ['Lock-in 거리', 'd_lock', '3.0 m', '본 연구 설정'],
    ['재평가 주기', 'Δt_re', '3.0 s', '본 연구 설정'],
]
t3 = doc.add_table(rows=len(gate_data), cols=4)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(gate_data):
    for j, cell_text in enumerate(row_data):
        cell = t3.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('게이트 선택 효용함수:').bold = True
p = doc.add_paragraph()
run = p.add_run('U(i) = β_dist × dᵢ + β_queue × qᵢ')
run.font.name = 'Cambria Math'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('dᵢ : 보행자에서 게이트 i까지의 유클리드 거리 (m)', style='List Bullet')
doc.add_paragraph('qᵢ : 게이트 i에 배정된 현재 대기 인원 (명)', style='List Bullet')

doc.add_paragraph(
    'Haghani & Sarvi(2017)는 멜버른 주요 철도역에서 승객 면접 조사(stated preference) 및 '
    '현장 관측(revealed preference) 데이터를 기반으로 혼합 네스티드 로짓 모델(Mixed Nested '
    'Logit Model)을 추정하였다. 출구 선택에 영향을 미치는 주요 요인으로 공간적 거리(distance)와 '
    '혼잡도(congestion/queue)를 확인하였으며, 본 연구의 β_dist = -0.25, β_queue = -0.30은 '
    '해당 연구에서 보고된 파라미터 범위(β_dist: -0.21 ~ -0.31, β_queue: -0.14 ~ -0.60) 내의 값이다.'
)

p = doc.add_paragraph()
p.add_run('다단계 의사결정 모델:').bold = True
doc.add_paragraph(
    '게이트 선택 후 비현실적인 핑퐁 효과(반복적 경로 변경)를 방지하기 위해 다음 세 가지 '
    '메커니즘을 적용한다.'
)
doc.add_paragraph(
    '관성(Inertia): 새 게이트의 효용이 현재 게이트 효용 + 전환 비용(C_switch = 1.5)을 '
    '초과해야 경로 변경을 실행한다.', style='List Number')
doc.add_paragraph(
    'Lock-in: 게이트까지 거리가 3.0m 이내이면 경로 변경을 비활성화한다.',
    style='List Number')
doc.add_paragraph(
    '재평가 주기: 경로 재평가를 3.0초 간격으로 제한하여 과도한 경로 변경을 방지한다.',
    style='List Number')

doc.add_heading('3.4 서비스 시간 파라미터', level=2)

svc_data = [
    ['유형', '분포', '파라미터', '범위', '비고'],
    ['태그 (교통카드)', '로그정규분포', 'μ=0.35, σ=0.35', '0.8~5.0초', '접근+태핑+통과'],
    ['태그리스', '-', '~0초', '0~0.3초', '무정지 통과'],
]
t4 = doc.add_table(rows=len(svc_data), cols=5)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(svc_data):
    for j, cell_text in enumerate(row_data):
        cell = t4.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    '태그 사용자의 서비스 시간은 로그정규분포로 모델링하였다. 로그정규분포는 서비스 시간의 '
    '비대칭성(대부분 빠르게 통과하나, 소수가 카드 인식 실패 등으로 지체)을 반영하는 데 적합하다. '
    '하한(0.8초)은 카드 태핑의 최소 물리적 시간, 상한(5.0초)은 카드 인식 실패나 잔액 부족 등 '
    '이상 상황을 반영한다.'
)

doc.add_heading('3.5 도착 모델 파라미터', level=2)

arr_data = [
    ['파라미터', '설정값', '근거'],
    ['열차 도착 간격', '180초', '서울 2호선 피크시 배차 간격 (~2.5~3분)'],
    ['1회 하차 인원', 'Poisson(λ=40)', '서쪽 계단 이용분 추정'],
    ['계단 진입 분산', 'N(7.5, 3.75)초', '15초에 걸쳐 분산 도착'],
]
t5 = doc.add_table(rows=len(arr_data), cols=3)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(arr_data):
    for j, cell_text in enumerate(row_data):
        cell = t5.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ============================================================
# 참고문헌
# ============================================================
doc.add_heading('참고문헌', level=1)

refs = [
    'Haghani, M. & Sarvi, M. (2017). Stated and revealed exit choices of pedestrian crowd evacuees. Transportation Research Part B: Methodological, 106, 410-427.',
    'Helbing, D. & Molnar, P. (1995). Social Force Model for Pedestrian Dynamics. Physical Review E, 51(5), 4282-4286.',
    'Kemloh Wagoum, A.U., Chraibi, M., & Zhang, J. (2015). JuPedSim: an open framework for simulating and analyzing the dynamics of pedestrians. 3rd Conference of Transportation Research Group of India.',
    'Kretz, T. (2015). On Oscillations in the Social Force Model. Physica A, 438, 272-285.',
    'Peng, J., Wei, Z., Li, J., Guo, X., & Wang, S. (2024). Passenger Flow Bottleneck Decongestion in Subway Stations: A Simulation Study. SIMULATION (SAGE).',
    'Rzezonka, J., Chraibi, M., Seyfried, A., Hein, B., & Schadschneider, A. (2022). An attempt to distinguish physical and socio-psychological influences on pedestrian bottleneck. Royal Society Open Science, 9(6), 211822.',
    'Seyfried, A., Steffen, B., Klingsch, W., & Boltes, M. (2005). The fundamental diagram of pedestrian movement revisited. Journal of Statistical Mechanics, P10002.',
    'Seyfried, A., Boltes, M., Kahler, J., Klingsch, W., Portz, A., Rupprecht, T., Schadschneider, A., Steffen, B., & Winkens, A. (2010). Enhanced empirical data for the fundamental diagram and the flow through bottlenecks. In: Pedestrian and Evacuation Dynamics 2008, Springer.',
    'Tordeux, A., Chraibi, M., & Seyfried, A. (2016). Collision-Free Speed Model for Pedestrian Dynamics. In: Traffic and Granular Flow \'15, pp. 225-232. Springer.',
    'Weidmann, U. (1993). Transporttechnik der Fussgaenger. Schriftenreihe des IVT Nr. 90, ETH Zuerich.',
    'Xu, Q., Chraibi, M., Tordeux, A., & Zhang, J. (2019). Generalized collision-free velocity model for pedestrian dynamics. Physica A, 535, 122521.',
    'Zhang, S., Chraibi, M., Zhang, J., Li, H., & Song, W. (2019). Validation of Collision-Free Speed Model in Bottlenecks. 2019 9th International Conference on Fire Science and Fire Protection Engineering (ICFSFPE), IEEE.',
    'Zhang, S., Zhang, J., Chraibi, M., & Song, W. (2021). A speed-based model for crowd simulation considering walking preferences. Communications in Nonlinear Science and Numerical Simulation, 95, 105624.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)

# 저장
output_path = r'C:\Users\aaron\tagless\docs\모델_선정_보고서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
