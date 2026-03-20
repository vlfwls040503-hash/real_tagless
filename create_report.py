"""시뮬레이션 파라미터 및 방법론 정리 Word 문서 생성"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pathlib

doc = Document()

# 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.5

# ── 제목 ──
title = doc.add_heading('성수역 태그리스 게이트 분리 배치 최적화', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('시뮬레이션 파라미터 및 방법론 정리 (v5)')
doc.add_paragraph('작성일: 2026-03-20')
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════
# 1. 시뮬레이션 개요
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. 시뮬레이션 개요', level=1)
doc.add_paragraph(
    '본 시뮬레이션은 서울 지하철 2호선 성수역 2F 대합실(서쪽 50m 구간)의 '
    '보행자 흐름을 재현한다. JuPedSim의 Collision-Free Speed Model V2 (CFSM V2)를 '
    '사용하며, 개별 보행자의 게이트 선택 의사결정과 서비스 시간을 모델링한다.'
)
doc.add_paragraph('')

t = doc.add_table(rows=7, cols=2)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('항목', '값'),
    ('시뮬레이션 도구', 'JuPedSim (Python)'),
    ('보행 모델', 'CollisionFreeSpeedModelV2 (CFSM V2)'),
    ('시뮬레이션 시간', '300초 (5분)'),
    ('시간 간격 (dt)', '0.05초'),
    ('대합실 크기', '50m × 25m (서쪽 절반)'),
    ('게이트 수', '7개'),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v

# ══════════════════════════════════════════════════════════════
# 2. 기하구조
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. 기하구조 (seongsu_west.py)', level=1)
doc.add_paragraph(
    'AnyLogic 도면 기반으로 좌표를 추출하여 JuPedSim Shapely 기하구조로 변환하였다. '
    '게이트 배리어는 시뮬레이션 기하구조에서 제외하고 시각화에서만 표시한다. '
    '이는 CFSM의 벽 반발력이 모든 기하구조에 동일하게 적용되어, '
    '좁은 게이트 통로(0.55m)에서 보행자가 벽에 밀착하는 비현실적 행동을 방지하기 위함이다.'
)

t = doc.add_table(rows=9, cols=3)
t.style = 'Light Grid Accent 1'
geo_data = [
    ('요소', '값', '비고'),
    ('대합실', '50m × 25m', 'x=동서, y=남북'),
    ('게이트 x좌표', '12.0m', '배리어 시작'),
    ('게이트 통과 길이', '1.5m', 'x방향'),
    ('게이트 통로 폭', '0.55m', 'y방향, 실측'),
    ('게이트 칸막이 폭', '0.30m', 'y방향'),
    ('계단 위치', 'x=3.0m', '상단: y=15~18, 하단: y=8~11'),
    ('출구 위치', 'x=26~29m', '상단: y=24, 하단: y=3'),
    ('비통행 구조물', 'x=30~48m', '상단: y=16~24, 하단: y=3~11'),
]
for i, row_data in enumerate(geo_data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

# ══════════════════════════════════════════════════════════════
# 3. CFSM V2 파라미터
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. 보행 모델 파라미터 (CFSM V2)', level=1)
doc.add_paragraph(
    'Collision-Free Speed Model V2는 속도 기반 미시적 보행 모델로, '
    '에이전트별 파라미터를 개별 설정할 수 있다. '
    '원 논문: Tordeux et al. (2015) "Collision-Free Speed Model for Pedestrian Dynamics" '
    '(arXiv:1512.05597)'
)

t = doc.add_table(rows=9, cols=4)
t.style = 'Light Grid Accent 1'
cfsm_data = [
    ('파라미터', '변수명', '값', '출처'),
    ('보행자 반경 (l)', 'radius', '0.225m', 'Weidmann, 1993'),
    ('희망속도 (v0)', 'desired_speed', 'N(1.34, 0.26) m/s', 'Weidmann, 1993'),
    ('시간간격 (T)', 'time_gap', '1.06s (대기열: 0.5s)', 'Tordeux et al., 2015'),
    ('타인 반발 강도 (a)', 'strength_neighbor', '8.0', 'JuPedSim default'),
    ('타인 반발 범위 (D)', 'range_neighbor', '0.1m', 'JuPedSim default'),
    ('벽 반발 강도', 'strength_geometry', '5.0', 'JuPedSim default'),
    ('벽 반발 범위', 'range_geometry', '0.02m', 'JuPedSim default'),
    ('대기열 time_gap', 'time_gap (queue)', '0.5s', '게이트 5m 이내 적용'),
]
for i, row_data in enumerate(cfsm_data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph(
    '* time_gap: 게이트 5m 이내에서는 0.5s로 감소시켜 대기열에서 '
    '바짝 붙어 서는 심리를 반영한다.'
)

# ══════════════════════════════════════════════════════════════
# 4. 도착 모델
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. 도착 모델 (열차 군집)', level=1)
doc.add_paragraph(
    '보행자를 균일하게 생성하지 않고, 열차 도착 시 한꺼번에 쏟아지는 '
    '군집(Platoon) 형태의 도착 모델을 사용한다.'
)

t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
arrival_data = [
    ('파라미터', '값', '비고'),
    ('열차 도착 간격', '180초', '2호선 피크시 약 2.5~3분'),
    ('하차 인원/회', 'Poisson(40)', '서쪽 계단 이용분'),
    ('계단 분산 시간', '|N(7.5, 3.75)|초', '정규분포 절대값'),
    ('첫 열차 도착', '5초', ''),
]
for i, row_data in enumerate(arrival_data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph(
    '큐잉 이론 검증: 게이트 1개당 최대 처리 용량 C = 1/Ts ≈ 0.66명/s. '
    '7개 게이트 총 용량 ≈ 4.6명/s. '
    '열차 1회 하차 40명이 15초에 분산 도착하면 피크 도착률 ≈ 2.7명/s < 4.6명/s → 해소 가능.'
)

# ══════════════════════════════════════════════════════════════
# 5. 게이트 선택 모델
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. 게이트 선택 모델 (Multinomial Logit)', level=1)

doc.add_heading('5.1 효용함수', level=2)
doc.add_paragraph('U(i) = beta_dist × dist(i) + beta_queue × queue(i)')
doc.add_paragraph('')

t = doc.add_table(rows=5, cols=4)
t.style = 'Light Grid Accent 1'
logit_data = [
    ('변수', '설명', '값', '출처'),
    ('beta_dist', '거리 민감도 계수', '-0.25', 'Haghani & Sarvi, 2016'),
    ('beta_queue', '대기열 민감도 계수', '-0.3', 'Haghani & Sarvi, 2016'),
    ('dist(i)', '보행자→게이트i 유클리드 거리', '(m)', ''),
    ('queue(i)', '게이트i에 배정된 대기 인원', '(명)', '배정 기반 전수 카운트'),
]
for i, row_data in enumerate(logit_data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

doc.add_heading('5.2 선택 확률', level=2)
doc.add_paragraph('P(i) = exp(U(i)) / sum(exp(U(j))),  j = 시야 내 게이트')
doc.add_paragraph('시야 반경(VISION_RADIUS) = 8.0m. 시야 밖 게이트는 선택 불가.')

doc.add_heading('5.3 밀도 측정 방식', level=2)
doc.add_paragraph(
    '기존: 게이트 앞 3m×1m 영역 내 물리적 위치 기반 카운트 → 대기열 과소 측정 문제.\n'
    '개선: 해당 게이트에 "배정된" 전체 미통과 에이전트 수 → 빈 게이트 정확 감지, 균등 분산 달성.'
)

doc.add_heading('5.4 선행연구 파라미터 범위', level=2)
doc.add_paragraph(
    'Haghani & Sarvi (2016) "How Simple Hypothetical-Choice Experiments Can Be '
    'Utilized to Learn Humans\' Navigational Escape Decisions":\n'
    '  - beta_dist: -0.21 ~ -0.31 (Revealed Choice 기준)\n'
    '  - beta_queue (congestion): -0.14 ~ -0.60\n'
    '  - Hypothetical에서는 대기열 가중치가 거리의 약 2배'
)

# ══════════════════════════════════════════════════════════════
# 6. 다단계 의사결정
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. 다단계 의사결정 모델', level=1)

doc.add_heading('6.1 관성 (Inertia / Switching Cost)', level=2)
doc.add_paragraph(
    '경로 변경 시 전환 비용(C_switch)을 적용한다:\n\n'
    '  경로 변경 조건: U(new) > U(current) + C_switch\n\n'
    '  C_switch = 1.5\n\n'
    '새 게이트의 효용이 현재 게이트 + 전환비용보다 높아야 변경. '
    '이는 0.5초마다 게이트를 바꾸는 핑퐁 효과(Oscillation)를 방지한다.'
)

doc.add_heading('6.2 Lock-in (경로 잠금)', level=2)
doc.add_paragraph(
    '게이트까지 거리 < 3.0m → 경로 변경 비활성화.\n'
    '실제로 줄에 합류한 후에는 옆줄이 빨라도 자기 줄을 유지하는 심리 반영.'
)

doc.add_heading('6.3 재평가 주기', level=2)
doc.add_paragraph(
    '3.0초 간격으로 재평가. Liao et al. (2017) 연구에서 전체 보행자의 '
    '0~7.5%만 경로를 변경한다는 결과에 기반.'
)

t = doc.add_table(rows=5, cols=3)
t.style = 'Light Grid Accent 1'
decision_data = [
    ('파라미터', '값', '설명'),
    ('C_switch', '1.5', '전환 비용 (관성)'),
    ('LOCK_IN_DISTANCE', '3.0m', '경로 잠금 거리'),
    ('REROUTE_INTERVAL', '3.0초', '재평가 주기'),
    ('참고', 'Liao et al., 2017', '경로 변경 비율 0~7.5%'),
]
for i, row_data in enumerate(decision_data):
    for j, val in enumerate(row_data):
        t.rows[i].cells[j].text = val

# ══════════════════════════════════════════════════════════════
# 7. 서비스 시간
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. 서비스 시간 모델 (게이트 통과)', level=1)
doc.add_paragraph(
    '서비스 시간: Lognormal(mu=0.35, sigma=0.35)\n'
    '  → 평균 약 1.5초, 범위 [0.8, 5.0]초 클리핑\n\n'
    '선행연구 참고:\n'
    '  - 베이징 지하철 NFC: 남성 무짐 2.91s, 여성 3.03s (Yang et al.)\n'
    '  - 한국 교통카드: 접촉 시간이 짧아 ~1.5s 적용\n\n'
    '게이트 구간(x=11.8~13.7) 진입 시:\n'
    '  1. desired_speed → 0.3 m/s (감속하며 카드 탭)\n'
    '  2. 서비스 시간 경과 후 → 원래 속도 복원\n\n'
    '대기열 제어:\n'
    '  - 게이트 사용 중(다른 에이전트 서비스 중) → 뒤에서 speed=0 정지 대기\n'
    '  - 게이트 비어있음 → 정상 속도로 진입'
)

# ══════════════════════════════════════════════════════════════
# 8. 3단계 경로 (Journey)
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. 3단계 경로 시스템', level=1)
doc.add_paragraph(
    'JuPedSim의 Journey/Stage 시스템을 활용한 3단계 경로:\n\n'
    '  1단계 - 접근 (Approach): Waypoint(x=8.0, gate_y), 반경 1.0m\n'
    '    → 게이트 y좌표로 정렬, 같은 줄로 합류\n\n'
    '  2단계 - 게이트 입구: Waypoint(x=12.0, gate_y), 반경 0.4m\n'
    '    → 좁은 반경으로 한 줄 대기 형성\n\n'
    '  3단계 - 게이트 통과 후: Waypoint(x=15.5, gate_y), 반경 1.0m\n'
    '    → 출구(상단/하단)로 이동, 퇴장'
)

# ══════════════════════════════════════════════════════════════
# 9. 주요 수정 이력
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. 주요 수정 이력 (v1 → v5)', level=1)

changes = [
    ('v1→v2', '기본 Queue stage 사용',
     '대기열이 로봇처럼 고정 위치에 서는 비현실적 행동 → Waypoint + 속도 제어로 전환'),
    ('v2→v3', 'Guide rail(물리적 가이드) 추가',
     '보행자가 가이드 레일에 갇혀 부자연스러운 움직임 → 가이드 레일 완전 제거'),
    ('v3→v4', 'CollisionFreeSpeedModel → V2',
     '에이전트별 파라미터 설정 필요 → CFSM V2로 전환, 선행연구 기반 캘리브레이션'),
    ('v4→v5(도착)', '균일 생성 → 열차 군집 모델',
     '생성률 > 처리용량으로 대기열 무한 증가 → 포아송 분포 군집 도착으로 변경'),
    ('v4→v5(선택)', '임의 파라미터 → 선행연구 기반',
     'beta_dist=-0.5, beta_queue=-2.0 → Haghani & Sarvi 기반 -0.25, -0.3'),
    ('v4→v5(핑퐁)', '매 0.5초 재평가 → 다단계 의사결정',
     '1,816회 경로변경 → 관성+Lock-in+3초 재평가로 4~38회'),
    ('v5(벽비비기)', '배리어를 기하구조에 포함',
     '벽 반발력이 동일하여 보행자가 게이트 벽에 밀착 → 배리어를 기하구조에서 제외, 코드로 제어'),
    ('v5(밀도)', '공간 기반 밀도 측정',
     '게이트 앞 3m만 측정하여 빈 게이트 미감지 → 배정 기반 전수 카운트로 변경'),
]

t = doc.add_table(rows=len(changes)+1, cols=3)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = '버전'
t.rows[0].cells[1].text = '변경 내용'
t.rows[0].cells[2].text = '이유'
for i, (ver, change, reason) in enumerate(changes):
    t.rows[i+1].cells[0].text = ver
    t.rows[i+1].cells[1].text = change
    t.rows[i+1].cells[2].text = reason

# ══════════════════════════════════════════════════════════════
# 10. 참고문헌
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. 참고문헌', level=1)
refs = [
    'Weidmann, U. (1993). Transporttechnik der Fussgänger. ETH Zürich.',
    'Tordeux, A. et al. (2015). Collision-Free Speed Model for Pedestrian Dynamics. arXiv:1512.05597.',
    'Haghani, M. & Sarvi, M. (2016). How Simple Hypothetical-Choice Experiments Can Be Utilized to Learn Humans\' Navigational Escape Decisions. PLoS ONE.',
    'Liao, W., Wagoum, A.K. & Bode, N.W.F. (2017). Route choice in pedestrians: determinants for initial choices and revising decisions. J. Royal Society Interface, 14(127).',
    'Yang, X. et al. (2025). Fuzzy-theory-based Social Force Model for Simulating Pedestrian Choice of Ticket Gates. Chinese Physics B.',
    'Helbing, D. & Molnár, P. (1995). Social Force Model for Pedestrian Dynamics. Physical Review E.',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Bullet')

# 저장
output_path = pathlib.Path(r'C:\Users\박필진\교통공학_졸업작품\tagless_project\output\시뮬레이션_파라미터_정리_v5.docx')
doc.save(str(output_path))
print(f"저장 완료: {output_path}")
